from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_DIR.mkdir(exist_ok=True)

BRAND_DARK = "064E4C"
BRAND_MID = "347F7A"
BRAND_LIGHT = "E7F1EF"
GRAY = "6B7280"


def set_cell_like_paragraph_box(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), BRAND_LIGHT)
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "14")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), BRAND_MID)
        borders.append(border)
    p_pr.append(borders)


def set_run(run, size, color, bold=False, spacing=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    if spacing is not None:
        spacing_element = OxmlElement("w:spacing")
        spacing_element.set(qn("w:val"), str(spacing))
        run._element.get_or_add_rPr().append(spacing_element)


def add_centered_text(doc, text, size, color, bold=False, before=0, after=0, spacing=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run(run, size, color, bold=bold, spacing=spacing)
    return paragraph


def build_document(ship, code, qr_file, output_file):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_after = Pt(6)
    logo_paragraph.add_run().add_picture(
        str(ROOT / "public" / "go-nordic-logo.png"),
        width=Cm(8.4),
    )

    add_centered_text(
        doc,
        "HANDOVERPRO",
        size=16,
        color=BRAND_DARK,
        bold=True,
        after=2,
        spacing=45,
    )
    add_centered_text(
        doc,
        ship,
        size=22,
        color=BRAND_MID,
        bold=True,
        after=9,
        spacing=80,
    )
    add_centered_text(
        doc,
        "Scan QR-koden",
        size=13,
        color=GRAY,
        bold=True,
        after=5,
    )

    qr_paragraph = doc.add_paragraph()
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_paragraph.paragraph_format.space_after = Pt(8)
    qr_paragraph.add_run().add_picture(str(qr_file), width=Cm(13.2))

    add_centered_text(
        doc,
        "Koden udfyldes automatisk - tryk OK",
        size=12,
        color=GRAY,
        after=7,
    )

    code_paragraph = doc.add_paragraph()
    code_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_paragraph.paragraph_format.space_before = Pt(0)
    code_paragraph.paragraph_format.space_after = Pt(10)
    set_cell_like_paragraph_box(code_paragraph)
    code_run = code_paragraph.add_run(code)
    set_run(
        code_run,
        size=30,
        color=BRAND_DARK,
        bold=True,
        spacing=100,
    )

    add_centered_text(
        doc,
        "Koden kan også fås hos din afdelingsleder.",
        size=10,
        color=GRAY,
        after=2,
    )
    add_centered_text(
        doc,
        "Enheden husker adgangen i 6 måneder.",
        size=9,
        color=GRAY,
        after=0,
    )

    doc.core_properties.title = f"HandoverPro {ship} QR-kode"
    doc.core_properties.subject = "QR-kode og fælles kabyskode"
    doc.core_properties.author = "HandoverPro"
    doc.save(output_file)


build_document(
    "NORDIC CROWN",
    "CROWN26",
    ROOT / "public" / "qr-crown.png",
    OUTPUT_DIR / "HandoverPro-Crown-QR.docx",
)
build_document(
    "NORDIC PEARL",
    "PEARL26",
    ROOT / "public" / "qr-pearl.png",
    OUTPUT_DIR / "HandoverPro-Pearl-QR.docx",
)
